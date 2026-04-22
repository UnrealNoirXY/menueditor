"""
╔══════════════════════════════════════════════════════════════════╗
║         NOIR MENU STUDIO  ·  v7.0  ·  Smart Edition          ║
╠══════════════════════════════════════════════════════════════════╣
║  Fix v7.0:                                                       ║
║  • Fix CRITICO: titoli doppi/tripli nel PDF (_split_bilingue)    ║
║  • Parser À La Carte riscritto: macchina a stati                 ║
║    → gestisce prezzo su riga separata + nomi bilingue nel nome   ║
║  • detect_pattern potenziato: analisi stili Word + €-sola        ║
║  • unisci_lingue smart merge (preserva EN già estratto)          ║
║  • Editor a pannelli: lista categorie ▸ form di editing           ║
║  • Mini-preview live del piatto nell'editor                      ║
║  • Layout Check: barre di riempimento per pagina logica          ║
║  • Template di stile: Classico / Moderno / Rustico               ║
║  • Retrocompatibilità v4/v5/v6                                   ║
╚══════════════════════════════════════════════════════════════════╝
"""

import streamlit as st
import docx
import re
import json
import base64
import io
import unicodedata
import pandas as pd

try:
    from weasyprint import HTML as WeasyHTML
    PDF_DISPONIBILE = True
except Exception:
    PDF_DISPONIBILE = False

st.set_page_config(page_title="Noir Menu Studio", layout="wide", page_icon="🍽️")


# ═══════════════════════════════════════════════════════════════════
# COSTANTI GLOBALI
# ═══════════════════════════════════════════════════════════════════

SEP_LINEA   = "📄 Linea + Titolo"
SEP_SPAZIO  = "➡️ Solo Spazio"
SEP_OPTIONS = [SEP_LINEA, SEP_SPAZIO]

LAYOUT_NUOVA    = "🆕 Nuova Pagina"
LAYOUT_CONTINUA = "➡️ Continua"
LAYOUT_STESSA   = "📄 Stessa Pagina"

TEMPLATE_CLASSICO = "🥂 Classico"
TEMPLATE_MODERNO  = "◼ Moderno"
TEMPLATE_RUSTICO  = "🌿 Rustico"
TEMPLATE_OPTIONS  = [TEMPLATE_CLASSICO, TEMPLATE_MODERNO, TEMPLATE_RUSTICO]

TEMPLATE_CSS = {
    TEMPLATE_CLASSICO: {
        'font_import': "@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,400;0,600;0,700;1,400;1,600&display=swap');",
        'font_family': "'Cormorant Garamond', Georgia, serif",
        'bg_color':    '#fdfaf6',
        'titolo_color':'#2b2b2b',
        'cat_color':   '#b58d3d',
        'nome_weight': '700',
        'desc_it_color':'#4a4a4a',
        'desc_en_color':'#a0a0a0',
        'prezzo_color': '#2b2b2b',
        'sep_color':   '#c8b99a',
        'footer_color':'#888',
    },
    TEMPLATE_MODERNO: {
        'font_import': "@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');",
        'font_family': "'Inter', 'Helvetica Neue', Arial, sans-serif",
        'bg_color':    '#ffffff',
        'titolo_color':'#111111',
        'cat_color':   '#111111',
        'nome_weight': '600',
        'desc_it_color':'#444444',
        'desc_en_color':'#888888',
        'prezzo_color': '#111111',
        'sep_color':   '#cccccc',
        'footer_color':'#aaaaaa',
    },
    TEMPLATE_RUSTICO: {
        'font_import': '',
        'font_family': "Georgia, 'Palatino Linotype', serif",
        'bg_color':    '#fdf8f0',
        'titolo_color':'#2c1810',
        'cat_color':   '#6b3d1e',
        'nome_weight': '700',
        'desc_it_color':'#4a3728',
        'desc_en_color':'#9a8070',
        'prezzo_color': '#2c1810',
        'sep_color':   '#c8a878',
        'footer_color':'#9a8070',
    },
}


# ═══════════════════════════════════════════════════════════════════
# CSS STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════

GLOBAL_CSS = """
<style>
.menu-banner {
    background: linear-gradient(135deg, #1a1205 0%, #3d2b0a 60%, #1a1205 100%);
    padding: 18px 32px 14px 32px; border-radius: 12px;
    margin-bottom: 20px; border: 1px solid #5a3e1b;
}
.menu-banner h1 { color: #d4a843; font-family: Georgia, serif; margin: 0; font-size: 1.9em; }
.menu-banner p  { color: #a08050; margin: 4px 0 0 0; font-size: 0.82em; letter-spacing: 0.05em; }
.page-map { background: #f7f3ec; border-radius: 8px; padding: 10px 14px;
            font-size: 0.82em; border: 1px solid #e0d0b0; margin-top: 6px; }
.page-map-row { margin-bottom: 4px; }
.page-num  { font-weight: 700; color: #b58d3d; }
.page-cats { color: #444; }
.pattern-badge {
    display: inline-block; background: #e8f4e8; border: 1px solid #a0c8a0;
    border-radius: 20px; padding: 3px 12px; font-size: 0.8em; color: #2d6a2d; margin-top: 6px;
}
.piatto-card {
    border: 1px solid #e0d0b0; border-radius: 8px; padding: 10px 16px;
    background: #fdfaf6; text-align: center; font-family: Georgia, serif; margin-top: 6px;
}
.piatto-card .nome { font-weight: 700; font-size: 1.05em; }
.piatto-card .desc { font-size: 0.88em; font-style: italic; color: #4a4a4a; margin: 2px 0; }
.piatto-card .desc-en { font-size: 0.82em; font-style: italic; color: #aaa; }
.piatto-card .prezzo { font-size: 0.95em; font-weight: 600; margin-top: 4px; }
.piatto-card .allerg { font-size: 0.75em; color: #b58d3d; font-style: italic; }
</style>
"""


# ═══════════════════════════════════════════════════════════════════
# UTILITY
# ═══════════════════════════════════════════════════════════════════

def get_image_base64(uploaded_file):
    if uploaded_file is not None:
        b64 = base64.b64encode(uploaded_file.getvalue()).decode()
        return f"data:{uploaded_file.type};base64,{b64}"
    return None


def _safe_str(val) -> str:
    if val is None:
        return ""
    try:
        if pd.isna(val):
            return ""
    except (TypeError, ValueError):
        pass
    s = str(val).strip()
    return "" if s.lower() == "nan" else s


def _safe_bool(val) -> bool:
    if isinstance(val, bool):
        return val
    try:
        if pd.isna(val):
            return False
    except (TypeError, ValueError):
        pass
    return str(val).strip().lower() in ("true", "1", "yes", "si", "sì")


def _safe_int(val, default=1) -> int:
    try:
        v = int(float(str(val)))
        return max(1, v)
    except Exception:
        return default


def _split_bilingue(testo: str) -> tuple:
    """
    Se il testo contiene ' / ' lo divide in (parte_IT, parte_EN).
    Es: "ANTIPASTI / Starters" → ("ANTIPASTI", "Starters")
    Altrimenti restituisce (testo, '').
    """
    if ' / ' in testo:
        parti = testo.split(' / ', 1)
        return parti[0].strip(), parti[1].strip()
    return testo.strip(), ''


def _normalizza_nome_fuzzy(n: str) -> str:
    """Normalizza per confronto fuzzy: minuscolo, senza accenti, senza punteggiatura."""
    n = n.lower().strip()
    n = unicodedata.normalize('NFD', n)
    n = ''.join(c for c in n if unicodedata.category(c) != 'Mn')
    n = re.sub(r'[^\w\s]', '', n)
    return re.sub(r'\s+', ' ', n).strip()


# ═══════════════════════════════════════════════════════════════════
# REGEX COMUNI
# ═══════════════════════════════════════════════════════════════════

_RE_EURO = re.compile(r'€')
_RE_PREZZO_FULL = re.compile(
    r'(€\s*\d[\d.,]*(?:\s*/\s*(?:per\s*100\s*g|l\'etto|etto|kg))?|\d[\d.,]*\s*€)',
    re.IGNORECASE,
)
_RE_PREZZO_SOLO = re.compile(r'^€\s*[\d.,]+\s*$')   # riga che è SOLO un prezzo
_RE_ALLERG_TONDE = re.compile(r'\((?:Allergen[^\)]*:|)\s*[\d,\s]+\)', re.IGNORECASE)
_RE_ALLERG_QUADRE = re.compile(r'\[\s*[\d,\s]+\s*\]')
_RE_SEP_PUNTINI = re.compile(r'\.{3,}')


def _normalizza_prezzo(raw: str) -> str:
    raw = raw.strip()
    num = re.sub(r'[€\s]', '', raw).replace(',', '.')
    return f"€ {num}"


def _estrai_allergeni_tonde(testo: str):
    trovati = []
    for m in _RE_ALLERG_TONDE.finditer(testo):
        pulito = re.sub(r'(Allergen[^\)]*:|Allergens\s*:)', '', m.group(0), flags=re.IGNORECASE)
        pulito = pulito.strip('() ').strip()
        trovati.append(pulito)
    clean = _RE_ALLERG_TONDE.sub('', testo).strip()
    return clean, ', '.join(trovati)


def _estrai_allergeni_quadre(testo: str):
    trovati = [m.group(0).strip('[] ').strip() for m in _RE_ALLERG_QUADRE.finditer(testo)]
    clean = _RE_ALLERG_QUADRE.sub('', testo).strip().strip('*').strip()
    return clean, ', '.join(trovati)


def _estrai_allergeni_auto(testo: str):
    clean, allerg = _estrai_allergeni_quadre(testo)
    if not allerg:
        clean, allerg = _estrai_allergeni_tonde(testo)
    return clean, allerg


# ═══════════════════════════════════════════════════════════════════
# RILEVAMENTO PATTERN — v7 potenziato
# ═══════════════════════════════════════════════════════════════════

def detect_pattern(doc) -> str:
    """
    Analizza stili Word + testo per rilevare il pattern:
    'alacarte' | 'bistrot' | 'pizza' | 'taglieri'
    """
    score = {'alacarte': 0, 'bistrot': 0, 'pizza': 0, 'taglieri': 0}

    paragrafi = list(doc.paragraphs)
    testi_raw = []
    for p in paragrafi:
        t = p.text.strip()
        if not t:
            continue
        if '\n' in t:
            testi_raw.extend(t.split('\n'))
        else:
            testi_raw.append(t)
    testi = [t.strip() for t in testi_raw if t.strip()][:60]

    # ── Segnali base dal testo ──────────────────────────────
    for t in testi:
        if re.search(r'.+\s*[—\-]\s*€\s*\d', t):
            score['alacarte'] += 2
        if t.startswith('-') and _RE_EURO.search(t):
            score['bistrot'] += 2
        if re.match(r'^\s*\(\d', t):
            score['bistrot'] += 1
        if _RE_SEP_PUNTINI.search(t) and _RE_EURO.search(t):
            score['bistrot'] += 2
        if re.search(r'taglier|apericena|platter', t, re.IGNORECASE):
            score['taglieri'] += 3
        if t.startswith('-') and not _RE_EURO.search(t):
            score['taglieri'] += 1

    # ── Segnale forte À La Carte v2: € sola su riga ────────
    righe_prezzo_sole = sum(1 for t in testi if _RE_PREZZO_SOLO.match(t))
    if righe_prezzo_sole >= 3:
        score['alacarte'] += righe_prezzo_sole * 3

    # ── Segnale forte À La Carte: stili Heading in Word ────
    heading_count = 0
    bold_no_euro = 0
    for p in paragrafi:
        t = p.text.strip()
        if not t:
            continue
        try:
            if 'Heading' in (p.style.name or ''):
                heading_count += 1
        except Exception:
            pass
        is_bold = bool(p.runs) and any(r.bold for r in p.runs if r.text.strip())
        if is_bold and not _RE_EURO.search(t):
            bold_no_euro += 1

    if heading_count >= 2:
        score['alacarte'] += heading_count * 2
    if bold_no_euro >= 4:
        score['alacarte'] += bold_no_euro

    # ── Segnale Pizza: coppie nome → riga con euro ──────────
    coppie_riga = sum(
        1 for i in range(len(testi) - 1)
        if not _RE_EURO.search(testi[i]) and _RE_EURO.search(testi[i+1])
        and len(testi[i].split()) <= 4
    )
    if coppie_riga >= 3:
        score['pizza'] += coppie_riga * 2

    para_multilinea_pizza = sum(
        1 for p in paragrafi
        if '\n' in p.text and _RE_EURO.search(p.text)
        and len(p.text.split('\n')) >= 2
        and not _RE_EURO.search(p.text.split('\n')[0])
        and len(p.text.split('\n')[0].split()) <= 4
    )
    if para_multilinea_pizza >= 3:
        score['pizza'] += para_multilinea_pizza * 3

    dash_euro = sum(
        1 for t in testi
        if re.search(r'[—]\s*€\s*\d', t)
        and not re.search(r'taglier|platter|apericena', t, re.IGNORECASE)
    )
    if dash_euro >= 4:
        score['alacarte'] += dash_euro * 2

    return max(score, key=score.get)


# ═══════════════════════════════════════════════════════════════════
# PARSER A — À La Carte  (v7: macchina a stati)
# Formato reale: Categoria (heading/caps) → Nome (bold) →
#                Desc IT (italic) → Desc EN (italic) → € XX (sola)
# ═══════════════════════════════════════════════════════════════════

def _is_categoria_alacarte(testo: str) -> bool:
    """Euristico per distinguere titoli categoria da nomi piatto."""
    if _RE_EURO.search(testo):
        return False
    if _RE_ALLERG_TONDE.search(testo):
        return False
    if len(testo) > 70:
        return False
    if testo.startswith('(') or testo.startswith('-'):
        return False
    # Parte IT tutta maiuscola → categoria
    it_part = testo.split('/')[0].strip().strip('*')
    if it_part and it_part == it_part.upper() and len(it_part) > 2:
        return True
    # Molto corto → probabile categoria
    if len(testo.split()) <= 4:
        return True
    return False


def parser_alacarte(doc) -> list:
    """
    Macchina a stati per menu À La Carte.
    Stati: IDLE → NOME → DESCRIZIONE → (prezzo chiude il piatto)

    Gestisce:
    - Prezzo su riga separata (formato reale Moresco)
    - Nome già bilingue "IT / EN (allerg)" → split automatico
    - Categorie bilingue "ANTIPASTI / Starters" → split automatico
    - Fallback: prezzo inline "Nome — €XX"
    """
    # Classifica ogni paragrafo
    paragrafi = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        try:
            is_heading = 'Heading' in (p.style.name or '')
        except Exception:
            is_heading = False
        is_bold = bool(p.runs) and any(r.bold for r in p.runs if r.text.strip())
        is_italic = bool(p.runs) and all(
            (r.italic or not r.text.strip()) for r in p.runs
        )
        ha_euro = bool(_RE_EURO.search(t))
        prezzo_solo = bool(_RE_PREZZO_SOLO.match(t))
        paragrafi.append({
            't': t, 'bold': is_bold, 'italic': is_italic,
            'heading': is_heading, 'ha_euro': ha_euro, 'prezzo_solo': prezzo_solo,
        })

    piatti = []
    cat_it = "Menu"
    cat_en = ""
    piatto = None   # dict del piatto in costruzione

    def _salva():
        nonlocal piatto
        if piatto and piatto.get('Nome'):
            piatto['Descrizione'] = re.sub(r'\s{2,}', ' ', piatto['Descrizione']).strip()
            piatti.append(piatto)
        piatto = None

    def _nuovo_piatto(nome_it, nome_en, allerg):
        nonlocal piatto
        _salva()
        piatto = {
            'Categoria': cat_it, 'Categoria EN': cat_en,
            'Nome': nome_it, 'Nome EN': nome_en,
            'Descrizione': '', 'Descrizione EN': '',
            'Prezzo': '', 'Allergeni': allerg,
        }

    desc_counter = 0   # conta le righe descrizione: 0=prima (IT), 1+=EN

    for para in paragrafi:
        t = para['t']

        # ── CASO 1: riga con SOLO il prezzo (€ XX) ──────────
        if para['prezzo_solo']:
            if piatto:
                piatto['Prezzo'] = _normalizza_prezzo(t)
                _salva()
                desc_counter = 0
            continue

        # ── CASO 2: prezzo inline su stessa riga del nome ───
        if para['ha_euro'] and not para['prezzo_solo']:
            m_p = _RE_PREZZO_FULL.search(t)
            if m_p:
                _salva()
                desc_counter = 0
                testo_clean, allerg = _estrai_allergeni_tonde(t)
                m2 = _RE_PREZZO_FULL.search(testo_clean)
                if m2:
                    prezzo = _normalizza_prezzo(m2.group(1))
                    prima = testo_clean[:m2.start()].strip()
                    dopo  = testo_clean[m2.end():].strip()
                    prima = re.sub(r'\s*[—\-\.…]+\s*$', '', prima).strip().strip('*').strip()
                    n_it, n_en = _split_bilingue(prima)
                    piatto = {
                        'Categoria': cat_it, 'Categoria EN': cat_en,
                        'Nome': n_it, 'Nome EN': n_en,
                        'Descrizione': dopo, 'Descrizione EN': '',
                        'Prezzo': prezzo, 'Allergeni': allerg,
                    }
                    _salva()
            continue

        # ── CASO 3: titolo categoria (heading Word o caps) ──
        if para['heading'] or (not para['ha_euro'] and _is_categoria_alacarte(t)):
            _salva()
            desc_counter = 0
            raw = t.strip('*').strip()
            cat_it, cat_en = _split_bilingue(raw)
            continue

        # ── CASO 4: nome piatto (bold, senza euro) ──────────
        if para['bold'] and not para['ha_euro']:
            desc_counter = 0
            testo_clean, allerg = _estrai_allergeni_tonde(t)
            n_it, n_en = _split_bilingue(testo_clean.strip('*').strip())
            _nuovo_piatto(n_it, n_en, allerg)
            continue

        # ── CASO 5: riga descrizione ─────────────────────────
        if piatto:
            testo_clean, allerg = _estrai_allergeni_tonde(t)
            if allerg:
                piatto['Allergeni'] += (', ' if piatto['Allergeni'] else '') + allerg
            if testo_clean:
                if desc_counter == 0:
                    piatto['Descrizione'] = (piatto['Descrizione'] + ' ' + testo_clean).strip()
                else:
                    piatto['Descrizione EN'] = (piatto['Descrizione EN'] + ' ' + testo_clean).strip()
                desc_counter += 1
            continue

    _salva()
    return piatti


# ═══════════════════════════════════════════════════════════════════
# PARSER B — Bistrot  (invariato v6)
# ═══════════════════════════════════════════════════════════════════

def parser_bistrot(doc) -> list:
    piatti = []
    categoria_attuale = "Menu"
    piatto_corrente = None

    for para in doc.paragraphs:
        testo_raw = para.text.strip()
        if not testo_raw:
            continue
        ha_euro = bool(_RE_EURO.search(testo_raw))

        if ha_euro:
            if piatto_corrente:
                piatti.append(piatto_corrente)
                piatto_corrente = None
            testo = re.sub(r'^[-•]\s*', '', testo_raw).strip()
            m_p = _RE_PREZZO_FULL.search(testo)
            if not m_p:
                continue
            prezzo = _normalizza_prezzo(m_p.group(1))
            prima = testo[:m_p.start()].strip()
            dopo  = testo[m_p.end():].strip()
            prima = _RE_SEP_PUNTINI.sub('', prima).strip()
            prima = re.sub(r'\s*[—\-]+\s*$', '', prima).strip()
            prima_clean, allerg = _estrai_allergeni_tonde(prima)
            nome = prima_clean.strip().strip('*').strip()
            desc = re.sub(r'^\*(.+)\*$', r'\1', dopo.strip().strip('*').strip())
            piatto_corrente = {
                'Categoria': categoria_attuale, 'Categoria EN': '',
                'Nome': nome, 'Nome EN': '',
                'Descrizione': desc, 'Descrizione EN': '',
                'Prezzo': prezzo, 'Allergeni': allerg,
            }
        else:
            testo = re.sub(r'^[-•]\s*', '', testo_raw).strip().strip('*').strip()
            if not testo:
                continue
            if _is_categoria_alacarte(testo):
                if piatto_corrente:
                    piatti.append(piatto_corrente)
                    piatto_corrente = None
                categoria_attuale = testo
            elif piatto_corrente:
                desc_extra, allerg_extra = _estrai_allergeni_tonde(testo)
                if allerg_extra:
                    piatto_corrente['Allergeni'] += (', ' if piatto_corrente['Allergeni'] else '') + allerg_extra
                if desc_extra:
                    piatto_corrente['Descrizione'] = (piatto_corrente['Descrizione'] + ' ' + desc_extra).strip()

    if piatto_corrente:
        piatti.append(piatto_corrente)
    for p in piatti:
        p['Descrizione'] = re.sub(r'\s{2,}', ' ', p['Descrizione']).strip()
    return piatti


# ═══════════════════════════════════════════════════════════════════
# PARSER C — Pizza  (invariato v6)
# ═══════════════════════════════════════════════════════════════════

def _parse_pizza_riga_con_euro(nome_pizza, riga_desc, categoria):
    riga_clean, allerg = _estrai_allergeni_tonde(riga_desc)
    if not allerg:
        riga_clean, allerg = _estrai_allergeni_quadre(riga_clean)
    m_p = _RE_PREZZO_FULL.search(riga_clean)
    if m_p:
        prezzo = _normalizza_prezzo(m_p.group(1))
        desc = (riga_clean[:m_p.start()] + ' ' + riga_clean[m_p.end():]).strip().strip(',').strip()
    else:
        prezzo = ''
        desc = riga_clean
    return {
        'Categoria': categoria, 'Categoria EN': '',
        'Nome': nome_pizza, 'Nome EN': '',
        'Descrizione': re.sub(r'\s{2,}', ' ', desc).strip(), 'Descrizione EN': '',
        'Prezzo': prezzo, 'Allergeni': allerg,
    }


def parser_pizza(doc) -> list:
    piatti = []
    categoria_attuale = 'Menu Pizza'
    righe_espanse = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if '\n' in t and _RE_EURO.search(t):
            linee = [l.strip() for l in t.split('\n') if l.strip()]
            if len(linee) >= 2 and not _RE_EURO.search(linee[0]):
                righe_espanse.append(linee[0])
                righe_espanse.append('\n'.join(linee[1:]))
                continue
        righe_espanse.append(t)

    i = 0
    while i < len(righe_espanse):
        testo = righe_espanse[i]
        ha_euro = bool(_RE_EURO.search(testo))
        if not ha_euro:
            prossima_ha_euro = (i + 1 < len(righe_espanse) and bool(_RE_EURO.search(righe_espanse[i+1])))
            if prossima_ha_euro and len(testo.split()) <= 5:
                i += 1
                piatti.append(_parse_pizza_riga_con_euro(testo.strip('*').strip(), righe_espanse[i], categoria_attuale))
            else:
                if len(testo.split()) <= 6:
                    categoria_attuale = testo.strip('*').strip()
            i += 1
        else:
            testo_clean, allerg = _estrai_allergeni_tonde(testo)
            if not allerg:
                testo_clean, allerg = _estrai_allergeni_quadre(testo_clean)
            m_p = _RE_PREZZO_FULL.search(testo_clean)
            if m_p:
                prezzo = _normalizza_prezzo(m_p.group(1))
                prima = testo_clean[:m_p.start()].strip()
                dopo  = testo_clean[m_p.end():].strip()
                if prima:
                    piatti.append({
                        'Categoria': categoria_attuale, 'Categoria EN': '',
                        'Nome': prima.strip().strip('*'), 'Nome EN': '',
                        'Descrizione': re.sub(r'\s{2,}', ' ', dopo).strip(), 'Descrizione EN': '',
                        'Prezzo': prezzo, 'Allergeni': allerg,
                    })
            i += 1
    return piatti


# ═══════════════════════════════════════════════════════════════════
# PARSER D — Taglieri  (invariato v6)
# ═══════════════════════════════════════════════════════════════════

def estrai_taglieri_word(file) -> list:
    doc = docx.Document(file)
    re_tag = re.compile(r'^(?:\d+\.\s*\.?\s*)?(?:-\s*)?(.+?)\s*[-–—]\s*(€\s*[\d,\.]+)', re.IGNORECASE)
    re_voce = re.compile(r'^[-•]\s+(.+)')
    rows = []
    ordine = 1
    paragrafi = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        is_italic = bool(para.runs) and all((r.italic or not r.text.strip()) for r in para.runs)
        is_italic = is_italic or (t.startswith('*') and t.endswith('*'))
        paragrafi.append({'testo': t, 'italic': is_italic})

    i = 0
    while i < len(paragrafi):
        p = paragrafi[i]
        testo = p['testo']
        if i == 0 and testo == testo.upper() and not _RE_EURO.search(testo):
            i += 1; continue
        if testo.startswith('(') and testo.endswith(')'):
            i += 1; continue

        if _RE_EURO.search(testo):
            m = re_tag.match(testo)
            if m:
                nome_raw = m.group(1).strip().strip('-').strip()
                prezzo = '€ ' + re.sub(r'[€\s]', '', m.group(2)).replace(',', '.')
                m_nick = re.match(r'^(.+?)\s*\(([^)]+)\)\s*$', nome_raw)
                if m_nick and not re.search(r'\d', m_nick.group(2)):
                    nome_it = m_nick.group(1).strip()
                    nome_en = m_nick.group(2).strip()
                else:
                    nome_it = nome_raw
                    nome_en = ''
                rows.append({
                    'Ordine': ordine, 'Tipo': 'Tagliere', 'Layout Pagina': LAYOUT_NUOVA,
                    'Visibile': True, 'Disponibile': True,
                    'Nome IT': nome_it, 'Nome EN': nome_en,
                    'Sottotitolo IT': '', 'Sottotitolo EN': '',
                    'Descrizione IT': '', 'Descrizione EN': '',
                    'Prezzo': prezzo, 'Allergeni': '', 'Note Chef': '', 'Scala': 1.0, 'Spazio Extra': 0.0,
                })
                ordine += 1
                i += 1
                if i < len(paragrafi):
                    np_ = paragrafi[i]
                    nt = np_['testo']
                    if np_['italic'] or (
                        not re_voce.match(nt) and not _RE_EURO.search(nt)
                        and not re_tag.match(nt) and not nt.startswith('(')
                    ):
                        rows[-1]['Sottotitolo IT'] = nt.strip('* ')
                        i += 1
                continue

        m_v = re_voce.match(testo)
        if m_v:
            contenuto = m_v.group(1).strip().strip('*').strip()
            contenuto, allerg = _estrai_allergeni_quadre(contenuto)
            if not allerg:
                contenuto, allerg = _estrai_allergeni_tonde(contenuto)
            if ':' in contenuto:
                nome_v, desc_v = contenuto.split(':', 1)
                nome_v = nome_v.strip().strip('*')
                desc_v = desc_v.strip().strip('*')
            else:
                nome_v = contenuto.strip('*').strip()
                desc_v = ''
            rows.append({
                'Ordine': ordine, 'Tipo': 'Voce', 'Layout Pagina': LAYOUT_NUOVA,
                'Visibile': True, 'Disponibile': True,
                'Nome IT': nome_v, 'Nome EN': '',
                'Sottotitolo IT': '', 'Sottotitolo EN': '',
                'Descrizione IT': desc_v, 'Descrizione EN': '',
                'Prezzo': '', 'Allergeni': allerg, 'Note Chef': '', 'Scala': 1.0, 'Spazio Extra': 0.0,
            })
            ordine += 1
        i += 1
    return rows


# ═══════════════════════════════════════════════════════════════════
# DISPATCHER
# ═══════════════════════════════════════════════════════════════════

PARSER_LABELS = {
    'alacarte': '🍽️ À La Carte',
    'bistrot':  '🥗 Bistrot / Trattoria',
    'pizza':    '🍕 Pizza',
    'taglieri': '🧀 Taglieri / Aperitivi',
}


def estrai_dati_word(file):
    doc = docx.Document(file)
    pattern = detect_pattern(doc)
    if pattern == 'pizza':
        return parser_pizza(doc), pattern
    elif pattern == 'bistrot':
        return parser_bistrot(doc), pattern
    elif pattern == 'taglieri':
        file.seek(0)
        return estrai_taglieri_word(file), pattern
    else:
        return parser_alacarte(doc), pattern


# ═══════════════════════════════════════════════════════════════════
# UNIFICATORE LINGUE — v7 smart merge
# ═══════════════════════════════════════════════════════════════════

def unisci_lingue(piatti_it: list, piatti_en: list) -> list:
    """
    Unisce IT ed EN con matching intelligente.
    Se IT contiene già dati EN (da _split_bilingue), li preserva.
    Se viene fornito file EN, lo usa per riempire i campi EN mancanti.
    """
    menu = []
    n = max(len(piatti_it), len(piatti_en))

    for i in range(n):
        it = piatti_it[i] if i < len(piatti_it) else {}
        en = piatti_en[i] if i < len(piatti_en) else {}

        # Preferisci dati EN già estratti dal parser IT (via _split_bilingue)
        nome_en = _safe_str(it.get('Nome EN', '')) or _safe_str(en.get('Nome', ''))
        cat_en  = _safe_str(it.get('Categoria EN', '')) or _safe_str(en.get('Categoria', 'Other'))
        desc_en = _safe_str(it.get('Descrizione EN', '')) or _safe_str(en.get('Descrizione', ''))

        # Sanity check: se EN risulta identico a IT (stesso testo), azzera
        if nome_en and nome_en == _safe_str(it.get('Nome', '')):
            nome_en = ''

        menu.append({
            'Ordine':         i + 1,
            'Pagina':         1,
            'Separatore':     SEP_LINEA,
            'Visibile':       True,
            'Disponibile':    True,
            'Categoria IT':   _safe_str(it.get('Categoria', 'Altro')),
            'Categoria EN':   cat_en,
            'Colore Cat.':    '#b58d3d',
            'Nome IT':        _safe_str(it.get('Nome', '')),
            'Nome EN':        nome_en,
            'Descrizione IT': _safe_str(it.get('Descrizione', '')).strip(),
            'Descrizione EN': desc_en.strip(),
            'Prezzo':         _safe_str(it.get('Prezzo', '') or en.get('Prezzo', '')),
            'Allergeni':      _safe_str(it.get('Allergeni', '') or en.get('Allergeni', '')),
            'Note Chef':      '',
            'Scala Piatto':   1.0,
            'Spazio Extra':   0.0,
        })
    return menu


# ═══════════════════════════════════════════════════════════════════
# NORMALIZZATORI COLONNE
# ═══════════════════════════════════════════════════════════════════

def _norm_separatore(v):
    s = _safe_str(v)
    if s in (LAYOUT_STESSA, SEP_LINEA):
        return SEP_LINEA
    if s in (LAYOUT_CONTINUA, SEP_SPAZIO):
        return SEP_SPAZIO
    return SEP_LINEA


def _norm_colore(v):
    s = _safe_str(v)
    return s if re.match(r'^#[0-9a-fA-F]{3,6}$', s) else '#b58d3d'


def _assicura_colonne_menu(df: pd.DataFrame) -> pd.DataFrame:
    if 'Layout Pagina' in df.columns and 'Pagina' not in df.columns:
        df['Pagina'] = 1
        if 'Separatore' not in df.columns:
            df['Separatore'] = df['Layout Pagina'].apply(_norm_separatore)
        df.drop(columns=['Layout Pagina'], inplace=True, errors='ignore')
    if 'Nuova Pagina Dopo' in df.columns and 'Pagina' not in df.columns:
        df['Pagina'] = 1
        df.drop(columns=['Nuova Pagina Dopo'], inplace=True, errors='ignore')

    defaults = {
        'Ordine':       lambda n: list(range(1, n + 1)),
        'Pagina':       1, 'Separatore': SEP_LINEA,
        'Visibile':     True, 'Disponibile': True,
        'Forza Salto Pagina': False,
        'Categoria EN': '', 'Colore Cat.': '#b58d3d',
        'Nome EN':      '', 'Descrizione EN': '',
        'Note Chef':    '', 'Scala Piatto': 1.0, 'Spazio Extra': 0.0,
    }
    for col, default in defaults.items():
        if col not in df.columns:
            df[col] = default(len(df)) if callable(default) else default

    df['Ordine']       = pd.to_numeric(df['Ordine'],       errors='coerce').fillna(0).astype(int)
    df['Pagina']       = df['Pagina'].apply(lambda x: _safe_int(x, 1))
    df['Scala Piatto'] = pd.to_numeric(df['Scala Piatto'], errors='coerce').fillna(1.0)
    df['Spazio Extra'] = pd.to_numeric(df['Spazio Extra'], errors='coerce').fillna(0.0)
    df['Visibile']     = df['Visibile'].apply(_safe_bool)
    df['Disponibile']  = df['Disponibile'].apply(_safe_bool)
    df['Forza Salto Pagina'] = df['Forza Salto Pagina'].apply(_safe_bool)
    df['Separatore']   = df['Separatore'].apply(_norm_separatore)
    df['Colore Cat.']  = df['Colore Cat.'].apply(_norm_colore)
    return df


def _assicura_colonne_aperitivi(df: pd.DataFrame) -> pd.DataFrame:
    if 'Nuova Pagina Dopo' in df.columns and 'Layout Pagina' not in df.columns:
        df['Layout Pagina'] = LAYOUT_NUOVA
        df.drop(columns=['Nuova Pagina Dopo'], inplace=True, errors='ignore')

    defaults = {
        'Ordine': lambda n: list(range(1, n + 1)),
        'Tipo': 'Voce', 'Layout Pagina': LAYOUT_NUOVA,
        'Visibile': True, 'Disponibile': True,
        'Nome IT': '', 'Nome EN': '', 'Sottotitolo IT': '', 'Sottotitolo EN': '',
        'Descrizione IT': '', 'Descrizione EN': '',
        'Prezzo': '', 'Allergeni': '', 'Note Chef': '', 'Scala': 1.0, 'Spazio Extra': 0.0,
    }
    for col, default in defaults.items():
        if col not in df.columns:
            df[col] = default(len(df)) if callable(default) else default

    def _norm_layout_aper(v):
        s = _safe_str(v)
        return s if s in (LAYOUT_NUOVA, LAYOUT_CONTINUA, LAYOUT_STESSA) else LAYOUT_NUOVA

    df['Ordine']       = pd.to_numeric(df['Ordine'],       errors='coerce').fillna(0).astype(int)
    df['Scala']        = pd.to_numeric(df['Scala'],        errors='coerce').fillna(1.0)
    df['Spazio Extra'] = pd.to_numeric(df['Spazio Extra'], errors='coerce').fillna(0.0)
    df['Visibile']     = df['Visibile'].apply(_safe_bool)
    df['Disponibile']  = df['Disponibile'].apply(_safe_bool)
    df['Layout Pagina']= df['Layout Pagina'].apply(_norm_layout_aper)
    return df


# ═══════════════════════════════════════════════════════════════════
# EXCEL HELPERS
# ═══════════════════════════════════════════════════════════════════

def df_to_excel_bytes(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as w:
        df.to_excel(w, index=False, sheet_name='Menu')
    return buf.getvalue()


def excel_bytes_to_df(raw):
    return pd.read_excel(io.BytesIO(raw), sheet_name='Menu', dtype=str)


def df_to_excel_bytes_aperitivi(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='openpyxl') as w:
        df.to_excel(w, index=False, sheet_name='Aperitivi')
    return buf.getvalue()


def excel_bytes_to_df_aperitivi(raw):
    return pd.read_excel(io.BytesIO(raw), sheet_name='Aperitivi', dtype=str)


# ═══════════════════════════════════════════════════════════════════
# LAYOUT CHECK — stima riempimento pagine
# ═══════════════════════════════════════════════════════════════════

_A4_H      = 297.0
_HEADER_P1 = 40.0
_HEADER_CN = 22.0
_FOOTER_H  = 36.0
_MARGIN_TOP= 10.0
_RIGA_BASE = 18.0


def _disp_foglio(primo: bool) -> float:
    return _A4_H - (_HEADER_P1 if primo else _HEADER_CN) - _FOOTER_H - _MARGIN_TOP


def _stima_piatto(row) -> float:
    try:
        scala = float(row.get('Scala Piatto', 1.0) or 1.0)
    except Exception:
        scala = 1.0
    try:
        extra = float(row.get('Spazio Extra', 0.0) or 0.0)
    except Exception:
        extra = 0.0
    desc_it = _safe_str(row.get('Descrizione IT', ''))
    desc_en = _safe_str(row.get('Descrizione EN', ''))
    righe = 1
    if desc_it: righe += max(1, len(desc_it) // 55)
    if desc_en: righe += max(1, len(desc_en) // 60)
    return (righe * _RIGA_BASE * scala) + (extra * _RIGA_BASE)


def render_layout_check(df: pd.DataFrame):
    """Barre di riempimento per pagina logica nella sidebar."""
    if df.empty or 'Pagina' not in df.columns:
        return
    pagine_usata = {}
    pagine_cats  = {}
    for _, row in df.iterrows():
        if not _safe_bool(row.get('Visibile', True)):
            continue
        pag = _safe_int(row.get('Pagina', 1))
        cat = _safe_str(row.get('Categoria IT', ''))
        pagine_usata[pag] = pagine_usata.get(pag, 0) + _stima_piatto(row)
        if cat and cat not in pagine_cats.get(pag, []):
            pagine_cats.setdefault(pag, []).append(cat)

    for pag in sorted(pagine_usata):
        usata = pagine_usata[pag]
        cap   = _disp_foglio(True)
        pct   = usata / cap
        bar   = min(pct * 100, 100)
        color = '#27ae60' if pct < 0.85 else '#f39c12' if pct < 1.0 else '#e74c3c'
        cats_str = ' · '.join(pagine_cats.get(pag, []))
        warn = '<div style="font-size:0.72em;color:#e74c3c;">⚠️ Eccede la pagina</div>' if pct > 1.0 else ''
        st.markdown(f"""
        <div style="margin-bottom:10px;">
            <div style="display:flex;justify-content:space-between;font-size:0.80em;margin-bottom:2px;">
                <span><b>P{pag}</b> {cats_str}</span>
                <span style="color:{color};font-weight:700;">{round(pct*100)}%</span>
            </div>
            <div style="background:#e8e8e8;border-radius:4px;height:7px;">
                <div style="background:{color};width:{bar:.1f}%;height:7px;border-radius:4px;transition:width 0.3s;"></div>
            </div>{warn}
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════
# GENERATORE HTML — Menu Principale  (v7: fix titoli doppi)
# ═══════════════════════════════════════════════════════════════════

def _build_css_menu(bg_b64, stile_sfondo, base_font_px, template_key) -> str:
    t = TEMPLATE_CSS.get(template_key, TEMPLATE_CSS[TEMPLATE_CLASSICO])

    if bg_b64:
        css_bg = (
            f"background-image:url('{bg_b64}');background-size:100% 100%;background-repeat:no-repeat;"
            if stile_sfondo == 'Adatta esattamente al foglio A4'
            else f"background-image:url('{bg_b64}');background-size:cover;background-position:center;"
        )
    else:
        css_bg = f"background-color:{t['bg_color']};"

    font_import = t.get('font_import', '')

    return f"""
{font_import}
@page {{ size:A4 portrait; margin:0; }}
*,*::before,*::after {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ background-color:#d0cfc9; font-family:{t['font_family']};
        margin:0; padding:24px 0; display:flex; flex-direction:column;
        align-items:center; gap:32px; }}
.foglio-a4 {{ width:210mm; height:297mm; font-size:{base_font_px}px;
              {css_bg} background-color:{t['bg_color']}; position:relative;
              page-break-after:always; break-after:page;
              box-shadow:0 6px 28px rgba(0,0,0,0.25); color:#2b2b2b; text-align:center; }}
.foglio-a4:last-child {{ page-break-after:auto; break-after:auto; }}
.content-area {{ padding:10mm 28mm 44mm 28mm; display:flex;
                 flex-direction:column; align-items:center; }}
.header {{ display:flex; flex-direction:column; align-items:center;
           width:100%; margin-bottom:2mm; }}
.logo-wrapper {{ line-height:0; margin-bottom:0.5em; }}
.titolo-menu {{ font-size:2.0em; font-weight:600; font-style:italic;
                color:{t['titolo_color']}; letter-spacing:0.04em; margin-bottom:2mm; line-height:1.2; }}
.titolo-categoria {{ font-size:1.55em; font-weight:600; font-style:italic;
                     color:{t['cat_color']}; display:inline-block;
                     border-bottom:1.5px solid {t['cat_color']}; padding-bottom:1mm;
                     margin-bottom:3mm; letter-spacing:0.03em; line-height:1.3; }}
.sep-categoria {{ width:75%; border:none; border-top:1px solid {t['sep_color']};
                  margin:4mm auto 2mm auto; opacity:0.6; }}
.piatti-area {{ width:100%; display:flex; flex-direction:column; align-items:center; }}
.blocco-piatto {{ max-width:85%; margin:0 auto 1.5em auto; line-height:1.5;
                  page-break-inside:avoid; break-inside:avoid;
                  text-align:center; width:100%; }}
.nome-piatto {{ font-size:1.18em; font-weight:{t['nome_weight']}; margin-bottom:0.15em;
                letter-spacing:0.02em; line-height:1.25; color:{t['titolo_color']}; }}
.badge-nd {{ display:inline-block; background:#c0392b; color:white;
             font-size:0.60em; font-weight:700; font-style:normal;
             padding:1px 5px; border-radius:3px; vertical-align:middle;
             margin-left:0.4em; letter-spacing:0.04em; }}
.allergeni {{ font-size:0.78em; color:{t['cat_color']}; font-style:italic; font-weight:400; }}
.desc-it {{ font-size:0.93em; font-style:italic; color:{t['desc_it_color']};
            margin:0.06em 0; line-height:1.42; }}
.desc-en {{ font-size:0.87em; font-style:italic; color:{t['desc_en_color']};
            margin:0.05em 0; line-height:1.38; }}
.prezzo {{ font-size:1.0em; font-weight:600; margin-top:0.28em;
           color:{t['prezzo_color']}; letter-spacing:0.01em; }}
.footer-area {{ position:absolute; bottom:12mm; left:28mm; right:28mm;
                text-align:center; border-top:0.5px solid {t['sep_color']};
                padding-top:2.5mm; }}
.footer {{ font-size:0.70em; font-style:italic; color:{t['footer_color']};
           margin:0; line-height:1.45; }}
@media print {{ body {{ background:none; padding:0; gap:0; }}
                .foglio-a4 {{ box-shadow:none; }} }}
"""


def _componi_titolo_cat(cat_it: str, cat_en: str, mostra_cat_en: bool) -> str:
    """
    FIX v7: evita titoli doppi/tripli.
    Se cat_it contiene già '/' (bilingue), non aggiungere cat_en.
    """
    if not mostra_cat_en or not cat_en:
        return cat_it
    if cat_en.lower() in ('other', 'altro', ''):
        return cat_it
    if '/' in cat_it:
        return cat_it   # già bilingue: non duplicare
    if cat_en == cat_it:
        return cat_it   # identici: non duplicare
    return f'{cat_it} / {cat_en}'


def genera_html(df: pd.DataFrame, logo_b64, bg_b64, stile_sfondo, titolo_menu,
                testo_footer, logo_size_px, base_font_px,
                mostra_nome_en=True, mostra_cat_en=True,
                template_key=TEMPLATE_CLASSICO,
                disabilita_autopaginazione=False) -> str:
    if 'Ordine' in df.columns:
        df = df.sort_values(['Pagina', 'Ordine']).reset_index(drop=True)

    css = _build_css_menu(bg_b64, stile_sfondo, base_font_px, template_key)
    logo_html   = (f'<div class="logo-wrapper"><img src="{logo_b64}" style="width:{logo_size_px}px;max-width:100%;"></div>'
                   if logo_b64 else '')
    titolo_html = f'<div class="titolo-menu">{titolo_menu}</div>' if titolo_menu else ''
    footer_block= (f'<div class="footer-area"><p class="footer">'
                   f'{testo_footer.replace(chr(10), "<br>")}</p></div>')

    parts = ['<!DOCTYPE html>\n<html>\n<head>', '<meta charset="UTF-8">',
             f'<style>{css}</style>', '</head>\n<body>']

    foglio_aperto   = False
    primo_foglio    = True
    primo_del_fog   = True
    usata           = 0.0
    cat_corrente    = None
    pagina_corrente = None
    col_corrente    = '#b58d3d'
    tit_corrente    = ''

    def apri_foglio(titolo_cat, colore, is_primo_assoluto, is_primo_gruppo):
        nonlocal foglio_aperto, primo_foglio, usata, primo_del_fog
        parts.append('<div class="foglio-a4"><div class="content-area"><div class="header">')
        parts.append(logo_html)
        if is_primo_assoluto and titolo_html:
            parts.append(titolo_html)
        if titolo_cat:
            parts.append(f'<div class="titolo-categoria" style="color:{colore};border-bottom-color:{colore};">'
                         f'{titolo_cat}</div>')
        parts.append('</div><div class="piatti-area">')
        foglio_aperto = True
        primo_del_fog = is_primo_gruppo
        usata = 0.0

    def chiudi_foglio():
        nonlocal foglio_aperto, primo_foglio
        parts.append('</div></div>')
        parts.append(footer_block)
        parts.append('</div>')
        foglio_aperto = False
        primo_foglio = False

    for _, row in df.iterrows():
        if not _safe_bool(row.get('Visibile', True)):
            continue

        pagina    = _safe_int(row.get('Pagina', 1))
        cat_it    = _safe_str(row.get('Categoria IT', ''))
        cat_en    = _safe_str(row.get('Categoria EN', ''))
        nome_it   = _safe_str(row.get('Nome IT', ''))
        nome_en   = _safe_str(row.get('Nome EN', ''))
        desc_it   = _safe_str(row.get('Descrizione IT', ''))
        desc_en   = _safe_str(row.get('Descrizione EN', ''))
        prezzo    = _safe_str(row.get('Prezzo', ''))
        allergeni = _safe_str(row.get('Allergeni', ''))
        disponibile = _safe_bool(row.get('Disponibile', True))
        separatore  = _norm_separatore(row.get('Separatore', SEP_LINEA))
        colore      = _norm_colore(row.get('Colore Cat.', '#b58d3d'))

        if not cat_it:
            continue

        # ── FIX CRITICO: costruzione titolo categoria senza duplicati ──
        titolo_cat = _componi_titolo_cat(cat_it, cat_en, mostra_cat_en)
        forza_salto = _safe_bool(row.get('Forza Salto Pagina', False))

        # Regola A: Cambio Pagina Manuale (Pagina != pagina_corrente)
        # Regola B: Forza Salto Explicit (forza_salto == True)
        if (pagina != pagina_corrente) or forza_salto:
            if foglio_aperto:
                chiudi_foglio()

            pagina_corrente = pagina

            # Gestione Titoli Categoria (Continua)
            # Se la categoria è la stessa della riga precedente (anche se abbiamo cambiato pagina),
            # aggiungiamo (Continua)
            tit_da_stampare = titolo_cat
            if cat_it == cat_corrente:
                tit_da_stampare = f"{titolo_cat} (Continua)"

            apri_foglio(tit_da_stampare, colore, primo_foglio, True)
            cat_corrente = cat_it
            col_corrente = colore
            tit_corrente = titolo_cat

        # Cambio categoria (stessa pagina)
        elif cat_it != cat_corrente:
            if not foglio_aperto:
                apri_foglio(titolo_cat, colore, primo_foglio, True)
            elif separatore == SEP_LINEA:
                parts.append('<hr class="sep-categoria">')
                parts.append(f'<div class="titolo-categoria" style="color:{colore};'
                              f'border-bottom-color:{colore};margin-bottom:3mm;">{titolo_cat}</div>')
            else:
                parts.append('<div style="margin-top:5mm;"></div>')
            cat_corrente = cat_it
            col_corrente = colore
            tit_corrente = titolo_cat

        # Regola C: Auto-paginazione Condizionata
        stima = _stima_piatto(row)
        if not disabilita_autopaginazione:
            if foglio_aperto and usata + stima > _disp_foglio(primo_del_fog):
                chiudi_foglio()

                # Gestione Titoli Categoria (Continua) in auto-paginazione
                tit_da_stampare = tit_corrente
                if cat_it == cat_corrente:
                    tit_da_stampare = f"{tit_corrente} (Continua)"

                apri_foglio(tit_da_stampare, col_corrente, False, False)

        # Rendering piatto
        try:
            scala = float(row.get('Scala Piatto', 1.0) or 1.0)
        except Exception:
            scala = 1.0
        try:
            extra = float(row.get('Spazio Extra', 0.0) or 0.0)
        except Exception:
            extra = 0.0

        nome_disp = f'{nome_it} / {nome_en}' if mostra_nome_en and nome_en else nome_it
        badge    = '<span class="badge-nd">NON DISPONIBILE</span>' if not disponibile else ''
        alrg_html= f' <span class="allergeni">({allergeni})</span>' if allergeni else ''
        dit_html = f'<div class="desc-it">{desc_it}</div>' if desc_it else ''
        den_html = f'<div class="desc-en">{desc_en}</div>' if (desc_en and mostra_nome_en) else ''
        prez_html= (f'<div class="prezzo"><s style="opacity:0.4;">{prezzo}</s></div>'
                    if not disponibile and prezzo
                    else (f'<div class="prezzo">{prezzo}</div>' if prezzo else ''))

        mb = round(1.5 * scala + extra, 3)
        parts.append(
            f'<div class="blocco-piatto" style="font-size:{scala}em;margin-bottom:{mb}em;">\n'
            f'  <div class="nome-piatto">{nome_disp}{badge}{alrg_html}</div>\n'
            f'  {dit_html}\n  {den_html}\n  {prez_html}\n</div>'
        )
        usata += stima

    if foglio_aperto:
        chiudi_foglio()
    parts.append('</body>\n</html>')
    return '\n'.join(parts)


# ═══════════════════════════════════════════════════════════════════
# GENERATORE HTML — Aperitivi  (invariato v6)
# ═══════════════════════════════════════════════════════════════════

def genera_html_aperitivi(df, logo_b64, bg_b64, stile_sfondo,
                           titolo_aperitivi, testo_footer, logo_size_px, base_font_px,
                           mostra_nome_en=True, pagina_unica=False) -> str:
    if 'Ordine' in df.columns:
        df = df.sort_values('Ordine').reset_index(drop=True)

    css_bg = (
        f"background-image:url('{bg_b64}');" + (
            'background-size:100% 100%;background-repeat:no-repeat;'
            if stile_sfondo == 'Adatta esattamente al foglio A4'
            else 'background-size:cover;background-position:center;'
        )
    ) if bg_b64 else 'background-color:#fdfaf6;'

    logo_html    = (f'<div class="logo-wrapper"><img src="{logo_b64}" style="width:{logo_size_px}px;max-width:100%;"></div>'
                    if logo_b64 else '')
    footer_block = (f'<div class="footer-area"><p class="footer">'
                    f'{testo_footer.replace(chr(10),"<br>")}</p></div>')

    m_tit = re.match(r'^(.+?)\s*\(([^)]+)\)\s*$', titolo_aperitivi)
    tit_p = m_tit.group(1).strip() if m_tit else titolo_aperitivi
    tit_s = m_tit.group(2).strip() if m_tit else ''

    css = f"""
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,400;0,600;0,700;1,400;1,600&display=swap');
@page {{ size:A4 portrait; margin:0; }}
*,*::before,*::after {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ background-color:#d0cfc9; font-family:'Cormorant Garamond',Georgia,serif;
        margin:0; padding:24px 0; display:flex; flex-direction:column; align-items:center; gap:32px; }}
.foglio-a4 {{ width:210mm; height:297mm; font-size:{base_font_px}px; {css_bg}
              background-color:white; position:relative; page-break-after:always; break-after:page;
              box-shadow:0 6px 28px rgba(0,0,0,0.25); color:#2b2b2b; }}
.foglio-a4:last-child {{ page-break-after:auto; break-after:auto; }}
.content-area {{ padding:10mm 26mm 44mm 26mm; display:flex; flex-direction:column; align-items:center; }}
.header {{ display:flex; flex-direction:column; align-items:center; width:100%; margin-bottom:3mm; }}
.logo-wrapper {{ line-height:0; margin-bottom:0.5em; }}
.titolo-aper {{ font-size:1.65em; font-weight:600; font-style:italic; color:#2b2b2b;
                letter-spacing:0.04em; margin-bottom:0.5mm; line-height:1.2; text-align:center; }}
.titolo-aper-sub {{ font-size:0.80em; font-style:italic; color:#999; margin-bottom:2.5mm; text-align:center; }}
.separatore {{ width:80%; border:none; border-top:1px solid #c8b99a; margin:0 auto 4mm auto; }}
.piatti-area {{ width:100%; display:flex; flex-direction:column; }}
.tagliere-block {{ width:100%; margin-bottom:4.5mm; page-break-inside:avoid; break-inside:avoid; }}
.tag-header-row {{ display:flex; align-items:baseline; border-bottom:1px solid #c8b99a;
                   padding-bottom:1.2mm; margin-bottom:1.5mm; gap:0.4em; }}
.tag-nome {{ font-size:1.15em; font-weight:700; color:#2b2b2b; flex-shrink:0; }}
.tag-nickname {{ font-size:0.78em; font-style:italic; color:#aaa; flex:1; }}
.tag-prezzo {{ font-size:1.12em; font-weight:700; color:#b58d3d; white-space:nowrap; flex-shrink:0; }}
.tag-sottotitolo {{ font-size:0.84em; font-style:italic; color:#6a6a6a; margin-bottom:1.8mm; text-align:center; }}
.voce-item {{ font-size:0.91em; line-height:1.48; padding:0.12em 0 0.12em 0.9em;
              border-left:2px solid #e8dcc8; margin-bottom:0.45em; }}
.voce-nome {{ font-weight:700; color:#2b2b2b; }}
.voce-sep {{ color:#b58d3d; margin:0 0.2em; }}
.voce-desc-it {{ font-style:italic; color:#4a4a4a; }}
.voce-desc-en {{ font-style:italic; color:#a8a8a8; font-size:0.90em; }}
.voce-allergeni {{ font-size:0.78em; color:#b58d3d; font-style:italic; margin-left:0.25em; }}
.badge-nd {{ display:inline-block; background:#c0392b; color:white; font-size:0.60em;
             font-weight:700; padding:1px 4px; border-radius:3px; vertical-align:middle; }}
.footer-area {{ position:absolute; bottom:12mm; left:26mm; right:26mm;
                text-align:center; border-top:0.5px solid #c8b99a; padding-top:2.5mm; }}
.footer {{ font-size:0.70em; font-style:italic; color:#888; margin:0; line-height:1.45; }}
@media print {{ body {{ background:none; padding:0; gap:0; }} .foglio-a4 {{ box-shadow:none; }} }}
"""
    parts = ['<!DOCTYPE html>\n<html>\n<head>', '<meta charset="UTF-8">',
             f'<style>{css}</style>', '</head>\n<body>']

    _HF=297.0; _HP=35.0; _HC=22.0; _FP=44.0; _MA=10.0; _HT=22.0; _HV=7.0
    def _disp_a(primo): return _HF - (_HP if primo else _HC) - _FP - _MA

    foglio_aperto=False; primo_foglio=True; tagliere_aperto=False; usata_a=0.0; primo_fog_a=True

    def _apri_a():
        nonlocal foglio_aperto, primo_foglio, usata_a, primo_fog_a
        parts.append('<div class="foglio-a4"><div class="content-area"><div class="header">')
        parts.append(logo_html)
        if primo_foglio and titolo_aperitivi:
            parts.append(f'<div class="titolo-aper">{tit_p}</div>')
            if tit_s: parts.append(f'<div class="titolo-aper-sub">({tit_s})</div>')
            primo_foglio = False
        parts.append('</div><hr class="separatore"><div class="piatti-area">')
        foglio_aperto=True; usata_a=0.0

    def _chiudi_a():
        nonlocal foglio_aperto, tagliere_aperto, primo_fog_a
        if tagliere_aperto:
            parts.append('</div>'); tagliere_aperto=False
        parts.append('</div></div>'); parts.append(footer_block); parts.append('</div>')
        foglio_aperto=False; primo_fog_a=False

    _apri_a()
    for _, row in df.iterrows():
        if not _safe_bool(row.get('Visibile', True)): continue
        tipo=_safe_str(row.get('Tipo','Voce')); nome_it=_safe_str(row.get('Nome IT'))
        nome_en=_safe_str(row.get('Nome EN')); sotto_it=_safe_str(row.get('Sottotitolo IT'))
        sotto_en=_safe_str(row.get('Sottotitolo EN')); desc_it=_safe_str(row.get('Descrizione IT'))
        desc_en=_safe_str(row.get('Descrizione EN')); prezzo=_safe_str(row.get('Prezzo'))
        allergeni=_safe_str(row.get('Allergeni')); disponibile=_safe_bool(row.get('Disponibile',True))
        layout_a=_safe_str(row.get('Layout Pagina',LAYOUT_NUOVA))
        try: scala=float(row.get('Scala',1.0) or 1.0)
        except: scala=1.0
        try: extra=float(row.get('Spazio Extra',0.0) or 0.0)
        except: extra=0.0
        if not nome_it: continue
        stima_a=(_HT if tipo=='Tagliere' else _HV)*scala+extra; disp_a=_disp_a(primo_fog_a)
        if not pagina_unica and foglio_aperto and usata_a+stima_a>disp_a:
            _chiudi_a(); _apri_a()
        if tipo=='Tagliere':
            if tagliere_aperto: parts.append('</div>'); tagliere_aperto=False
            if not pagina_unica and layout_a==LAYOUT_NUOVA and usata_a>0:
                _chiudi_a(); _apri_a()
            badge='<span class="badge-nd">N.D.</span>' if not disponibile else ''
            nick=f'<span class="tag-nickname">({nome_en})</span>' if nome_en else ''
            mb=round(4.5*scala+extra,3)
            parts.append(f'<div class="tagliere-block" style="margin-bottom:{mb}mm;">')
            parts.append(f'<div class="tag-header-row"><span class="tag-nome" style="font-size:{scala}em;">'
                         f'{nome_it}{badge}</span>{nick}<span class="tag-prezzo">{prezzo}</span></div>')
            if sotto_it: parts.append(f'<div class="tag-sottotitolo">{sotto_it}</div>')
            if sotto_en and mostra_nome_en: parts.append(f'<div class="tag-sottotitolo" style="color:#b0b0b0;font-size:0.88em;">{sotto_en}</div>')
            tagliere_aperto=True
        else:
            alrg=f'<span class="voce-allergeni">({allergeni})</span>' if allergeni else ''
            nd_v='<span class="badge-nd">N.D.</span>' if not disponibile else ''
            ndis=f'{nome_it} / {nome_en}' if mostra_nome_en and nome_en else nome_it
            dit=f'<span class="voce-sep">·</span><span class="voce-desc-it">{desc_it}</span>' if desc_it else ''
            den=f'<div class="voce-desc-en">{desc_en}</div>' if (desc_en and mostra_nome_en) else ''
            mb=round(0.45*scala+extra,3)
            parts.append(f'<div class="voce-item" style="font-size:{scala}em;margin-bottom:{mb}em;">'
                         f'<span class="voce-nome">{ndis}{nd_v}</span>{dit}{alrg}{den}</div>')
        usata_a+=stima_a
    if foglio_aperto: _chiudi_a()
    parts.append('</body>\n</html>')
    return '\n'.join(parts)


# ═══════════════════════════════════════════════════════════════════
# MAPPA PAGINE
# ═══════════════════════════════════════════════════════════════════

def build_mappa_pagine(df: pd.DataFrame) -> dict:
    if df.empty or 'Categoria IT' not in df.columns:
        return {}
    mappa = {}
    for _, row in df.iterrows():
        if not _safe_bool(row.get('Visibile', True)):
            continue
        pag = _safe_int(row.get('Pagina', 1))
        cat = _safe_str(row.get('Categoria IT', ''))
        if not cat:
            continue
        if pag not in mappa:
            mappa[pag] = []
        if cat not in mappa[pag]:
            mappa[pag].append(cat)
    return dict(sorted(mappa.items()))


def render_mappa_pagine(mappa: dict):
    if not mappa:
        st.caption("Nessun dato caricato.")
        return
    righe = []
    for pag, cats in mappa.items():
        cats_str = ' · '.join(cats)
        righe.append(f'<div class="page-map-row"><span class="page-num">📄 Pag. {pag}:</span> '
                     f'<span class="page-cats">{cats_str}</span></div>')
    st.markdown(f'<div class="page-map">{"".join(righe)}</div>', unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════
# MINI PREVIEW piatto (per editor pannello)
# ═══════════════════════════════════════════════════════════════════

def html_mini_preview(nome_it, nome_en, desc_it, desc_en, prezzo, allergeni,
                      colore='#b58d3d', mostra_en=True) -> str:
    nome_disp = f'{nome_it} / {nome_en}' if mostra_en and nome_en else nome_it
    al_html = f'<span style="font-size:0.75em;color:{colore};font-style:italic;"> ({allergeni})</span>' if allergeni else ''
    di_html = f'<div class="desc">{desc_it}</div>' if desc_it else ''
    de_html = f'<div class="desc-en">{desc_en}</div>' if (desc_en and mostra_en) else ''
    p_html  = f'<div class="prezzo">{prezzo}</div>' if prezzo else ''
    return f"""
    <div class="piatto-card">
        <div class="nome">{nome_disp}{al_html}</div>
        {di_html}{de_html}{p_html}
    </div>"""


# ═══════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown('## 🍽️ Menu Studio v7')

    with st.expander('🎨 Grafica & Template', expanded=True):
        template_sel     = st.selectbox('Template stile', TEMPLATE_OPTIONS, index=0)
        titolo_menu      = st.text_input('Titolo menu', 'Menu Alla Carta')
        titolo_aperitivi = st.text_input('Titolo aperitivi', 'Menu Apericena')
        logo_file        = st.file_uploader('Logo (PNG/JPG)', type=['png','jpg','jpeg'])
        logo_size        = st.slider('Dimensione Logo (px)', 50, 300, 130)
        bg_file          = st.file_uploader('Sfondo / Texture', type=['png','jpg','jpeg'])
        stile_sfondo     = st.radio('Adattamento sfondo',
                                    ['Riempi (Taglia i bordi)', 'Adatta esattamente al foglio A4'])
        testo_footer = st.text_area('Piè di pagina',
                                    'Il prodotto ha seguito la Catena del Freddo al fine di '
                                    "garantirne l'integrità e la sicurezza alimentare.")

    with st.expander('🌐 Lingue'):
        mostra_nome_en = st.toggle('Mostra nomi piatti in inglese', value=True)
        mostra_cat_en  = st.toggle('Mostra categorie in inglese',   value=True)

    with st.expander('📐 Testo & Impaginazione'):
        zoom_foglio = st.slider('Scala testi', 0.55, 1.20, 0.95, 0.05,
                                help='16px × zoom = font-size base del foglio A4')
        disabilita_autopaginazione = st.toggle('🛑 Disabilita auto-paginazione (Controllo 100% Manuale)', value=False)

    with st.expander('🍷 Aperitivi'):
        pagina_unica_aper = st.toggle('Tutto su un foglio A4', value=False)

    with st.expander('📁 Nomi File Export'):
        nome_file_menu = st.text_input('Nome file menu', 'Menu_Noir')
        nome_file_aper = st.text_input('Nome file aperitivi', 'Menu_Aperitivi')

    with st.expander('💾 Carica Progetto Menu'):
        progetto_caricato = st.file_uploader('Progetto (.json)', type=['json'], key='menu_json')
        progetto_xlsx     = st.file_uploader('Da Excel (.xlsx)',  type=['xlsx'], key='menu_xlsx')

    with st.expander('🍷 Carica Progetto Aperitivi'):
        aper_progetto_json = st.file_uploader('Aperitivi (.json)', type=['json'], key='aper_json')
        aper_progetto_xlsx = st.file_uploader('Aperitivi (.xlsx)', type=['xlsx'], key='aper_xlsx')

    st.divider()
    st.markdown('### 🗺️ Mappa Pagine')
    _mappa_placeholder = st.empty()

    st.divider()
    st.markdown('### 📊 Layout Check')
    st.caption('Stima riempimento per pagina logica.')
    _layout_check_placeholder = st.empty()


# ═══════════════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════════════

st.markdown(GLOBAL_CSS, unsafe_allow_html=True)
st.markdown(
    '<div class="menu-banner">'
    '<h1>🍽️ Noir Menu Studio</h1>'
    '<p>v7.0 · Smart Edition · Parser macchina-a-stati · Editor pannelli · Layout Check · Template</p>'
    '</div>', unsafe_allow_html=True,
)

tab_menu, tab_aperitivi = st.tabs(['🍕 Menu Ristorante / Pizzeria', '🍷 Menu Aperitivi / Taglieri'])


# ═══════════════════════════════════════════════════════════════════
# TAB 1 — MENU PRINCIPALE
# ═══════════════════════════════════════════════════════════════════

with tab_menu:

    # Session state init
    for k, v in [
        ('dati_menu', pd.DataFrame()), ('menu_json_id', None),
        ('menu_xlsx_id', None), ('menu_pattern', None), ('piatto_sel', None),
    ]:
        if k not in st.session_state:
            st.session_state[k] = v

    # ── Caricamento Word ────────────────────────────────────────
    st.markdown('### 📂 Caricamento Testi')
    col1, col2 = st.columns(2)
    with col1:
        file_it = st.file_uploader('🇮🇹 Word Italiano (.docx)', type=['docx'], key='pizza_it')
    with col2:
        file_en = st.file_uploader('🇬🇧 Word Inglese (.docx) — opzionale', type=['docx'], key='pizza_en')

    col_btn, col_info = st.columns([1, 3])
    with col_btn:
        btn_estrai = st.button('🔄 Estrai e Analizza', type='primary',
                               key='btn_pizza', use_container_width=True)

    if btn_estrai:
        if file_it:
            with st.spinner('Parser in esecuzione…'):
                piatti_it, pattern_it = estrai_dati_word(file_it)
                piatti_en, pattern_en = (estrai_dati_word(file_en) if file_en else ([], pattern_it))

            dati_uniti = unisci_lingue(piatti_it, piatti_en)
            df_new = pd.DataFrame(dati_uniti)
            st.session_state.dati_menu    = df_new
            st.session_state.menu_pattern = pattern_it
            st.session_state.menu_json_id = None
            st.session_state.menu_xlsx_id = None
            st.session_state.piatto_sel   = None

            n_it = len(piatti_it); n_en = len(piatti_en) if piatti_en else 0
            n_cat = df_new['Categoria IT'].nunique()
            with col_info:
                lbl = PARSER_LABELS.get(pattern_it, pattern_it)
                st.markdown(f'<span class="pattern-badge">✅ Pattern: {lbl}</span>',
                            unsafe_allow_html=True)
                if n_en and n_it != n_en:
                    st.warning(f'⚠️ {n_it} piatti IT vs {n_en} EN — verifica allineamento nell\'editor.')
                else:
                    st.success(f'✅ {n_it} piatti · {n_cat} categorie estratte.')
        else:
            st.error('⚠️ Carica almeno il file Word italiano.')

    # ── Carica progetto JSON ────────────────────────────────────
    if progetto_caricato is not None:
        fid = id(progetto_caricato)
        if fid != st.session_state.menu_json_id:
            try:
                st.session_state.dati_menu    = _assicura_colonne_menu(pd.DataFrame(json.load(progetto_caricato)))
                st.session_state.menu_json_id = fid
                st.session_state.piatto_sel   = None
                st.success('✅ Progetto JSON caricato.')
            except Exception as e:
                st.error(f'Errore JSON: {e}')

    if progetto_xlsx is not None:
        fid = id(progetto_xlsx)
        if fid != st.session_state.menu_xlsx_id:
            try:
                st.session_state.dati_menu    = _assicura_colonne_menu(excel_bytes_to_df(progetto_xlsx.getvalue()))
                st.session_state.menu_xlsx_id = fid
                st.session_state.piatto_sel   = None
                st.success('✅ Dati Excel caricati.')
            except Exception as e:
                st.error(f'Errore Excel: {e}')

    # ── Editor ─────────────────────────────────────────────────
    if not st.session_state.dati_menu.empty:
        df_base = _assicura_colonne_menu(st.session_state.dati_menu.copy())
        df_base = df_base.sort_values(['Pagina', 'Ordine']).reset_index(drop=True)
        # Rifletti nel session state con indice pulito
        st.session_state.dati_menu = df_base.copy()

        st.divider()

        # Metriche
        n_tot = len(df_base)
        n_vis = int(df_base['Visibile'].apply(_safe_bool).sum())
        n_nd  = int((~df_base['Disponibile'].apply(_safe_bool)).sum())
        n_cat = df_base['Categoria IT'].apply(_safe_str).nunique()
        n_pag = df_base['Pagina'].apply(lambda x: _safe_int(x,1)).nunique()
        mc1,mc2,mc3,mc4,mc5 = st.columns(5)
        mc1.metric('📋 Piatti',         n_tot)
        mc2.metric('👁️ Visibili',        n_vis)
        mc3.metric('🔴 Non dispon.',     n_nd)
        mc4.metric('📂 Categorie',       n_cat)
        mc5.metric('📄 Pagine logiche',  n_pag)

        # Aggiorna sidebar
        mappa = build_mappa_pagine(df_base)
        with _mappa_placeholder:
            render_mappa_pagine(mappa)
        with _layout_check_placeholder:
            render_layout_check(df_base)

        # ────────────────────────────────────────────────────────
        # PANNELLO: Organizzazione Pagine (Assegnazione Massiva)
        # ────────────────────────────────────────────────────────
        with st.expander("🏗️ Organizzazione Pagine (Assegnazione Massiva)", expanded=False):
            st.markdown("Sposta rapidamente tutte le portate di una categoria in una pagina specifica.")

            # Ottieni categorie uniche mantenendo l'ordine
            categorie_uniche = []
            for c in df_base['Categoria IT'].apply(_safe_str).tolist():
                if c and c not in categorie_uniche:
                    categorie_uniche.append(c)

            if categorie_uniche:
                c1, c2, c3 = st.columns([2, 1, 1])
                with c1:
                    cat_da_spostare = st.selectbox("Seleziona Categoria", categorie_uniche)
                with c2:
                    nuova_pag_target = st.number_input("Alla Pagina", min_value=1, value=1, step=1)
                with c3:
                    st.markdown("<div style='margin-top:28px;'></div>", unsafe_allow_html=True)
                    if st.button("Sposta Tutti", use_container_width=True):
                        mask = st.session_state.dati_menu['Categoria IT'].apply(_safe_str) == cat_da_spostare
                        st.session_state.dati_menu.loc[mask, 'Pagina'] = int(nuova_pag_target)
                        st.success(f"✅ Tutti i '{cat_da_spostare}' spostati a Pagina {nuova_pag_target}")
                        st.rerun()
            else:
                st.caption("Nessuna categoria trovata.")

        st.markdown('### ✏️ Editor Piatti')

        # ────────────────────────────────────────────────────────
        # PANNELLO: lista sinistra + form destra
        # ────────────────────────────────────────────────────────
        col_lista, col_form = st.columns([1, 2], gap="medium")

        with col_lista:
            st.markdown("**📋 Lista piatti**")

            # Bottone aggiungi nuovo piatto
            if st.button('➕ Nuovo Piatto', use_container_width=True, key='btn_nuovo'):
                cat_attuale = _safe_str(df_base.iloc[-1]['Categoria IT']) if not df_base.empty else 'Menu'
                new_ord = int(df_base['Ordine'].max()) + 1 if not df_base.empty else 1
                new_row = pd.DataFrame([{
                    'Ordine': new_ord, 'Pagina': 1, 'Separatore': SEP_LINEA,
                    'Visibile': True, 'Disponibile': True,
                    'Categoria IT': cat_attuale, 'Categoria EN': '',
                    'Colore Cat.': '#b58d3d', 'Nome IT': '(nuovo piatto)', 'Nome EN': '',
                    'Descrizione IT': '', 'Descrizione EN': '',
                    'Prezzo': '', 'Allergeni': '', 'Note Chef': '',
                    'Scala Piatto': 1.0, 'Spazio Extra': 0.0,
                }])
                st.session_state.dati_menu = pd.concat(
                    [st.session_state.dati_menu, new_row], ignore_index=True
                )
                st.session_state.piatto_sel = len(st.session_state.dati_menu) - 1
                st.rerun()

            st.markdown("---")

            # Raggruppa per categoria
            df_ord = st.session_state.dati_menu.copy()
            categorie = []
            for cat in df_ord['Categoria IT'].apply(_safe_str).tolist():
                if cat not in categorie:
                    categorie.append(cat)

            for cat in categorie:
                mask = df_ord['Categoria IT'].apply(_safe_str) == cat
                piatti_cat = df_ord[mask]
                n_c = len(piatti_cat)
                n_vis_c = piatti_cat['Visibile'].apply(_safe_bool).sum()

                with st.expander(f"📂 {cat}  ({n_vis_c}/{n_c})", expanded=True):
                    for idx, row_c in piatti_cat.iterrows():
                        nome_c = _safe_str(row_c.get('Nome IT')) or '(senza nome)'
                        is_avail = _safe_bool(row_c.get('Disponibile', True))
                        is_vis_r = _safe_bool(row_c.get('Visibile', True))
                        is_sel   = (st.session_state.piatto_sel == idx)

                        icona = '▶ ' if is_sel else ''
                        suf   = (' 🔴' if not is_avail else '') + (' 🚫' if not is_vis_r else '')
                        label = f"{icona}{nome_c}{suf}"

                        if st.button(label, key=f'sel_{idx}',
                                     type='primary' if is_sel else 'secondary',
                                     use_container_width=True):
                            st.session_state.piatto_sel = idx
                            st.rerun()

        # ── Form di editing ─────────────────────────────────────
        with col_form:
            sel = st.session_state.piatto_sel
            df_edit = st.session_state.dati_menu

            # Valida selezione
            if sel is None or sel not in df_edit.index:
                st.info("👈 Seleziona un piatto dalla lista per modificarlo, oppure premi ➕ per aggiungerne uno nuovo.")
            else:
                row = df_edit.loc[sel]
                nome_display = _safe_str(row.get('Nome IT', '')) or 'Nuovo Piatto'
                st.markdown(f"**✏️ Modifica: {nome_display}**")

                # Mini preview live (mostrata sopra le tab)
                colore_cat = _norm_colore(row.get('Colore Cat.', '#b58d3d'))
                preview_html = html_mini_preview(
                    _safe_str(row.get('Nome IT','')), _safe_str(row.get('Nome EN','')),
                    _safe_str(row.get('Descrizione IT','')), _safe_str(row.get('Descrizione EN','')),
                    _safe_str(row.get('Prezzo','')), _safe_str(row.get('Allergeni','')),
                    colore_cat, mostra_nome_en,
                )
                st.markdown(preview_html, unsafe_allow_html=True)

                tab_contenuto, tab_layout_ed, tab_avanzato = st.tabs(
                    ["📝 Contenuto", "📄 Layout & Pagina", "⚙️ Avanzato"]
                )

                # ── TAB CONTENUTO ─────────────────────────────
                with tab_contenuto:
                    with st.form(key=f'form_cont_{sel}'):
                        c1, c2 = st.columns(2)
                        with c1:
                            f_nome_it = st.text_input('🇮🇹 Nome IT',
                                                      value=_safe_str(row.get('Nome IT','')))
                        with c2:
                            f_nome_en = st.text_input('🇬🇧 Nome EN',
                                                      value=_safe_str(row.get('Nome EN','')))
                        f_desc_it = st.text_area('📝 Descrizione IT',
                                                 value=_safe_str(row.get('Descrizione IT','')),
                                                 height=80)
                        f_desc_en = st.text_area('📝 Descrizione EN',
                                                 value=_safe_str(row.get('Descrizione EN','')),
                                                 height=80)
                        c3, c4 = st.columns(2)
                        with c3:
                            f_prezzo  = st.text_input('💶 Prezzo',
                                                      value=_safe_str(row.get('Prezzo','')))
                        with c4:
                            f_allerg  = st.text_input('⚠️ Allergeni (es: 1, 7, 14)',
                                                      value=_safe_str(row.get('Allergeni','')))
                        c5, c6 = st.columns(2)
                        with c5:
                            f_vis  = st.checkbox('👁️ Visibile nel menu',
                                                 value=_safe_bool(row.get('Visibile', True)))
                        with c6:
                            f_disp = st.checkbox('✅ Disponibile',
                                                 value=_safe_bool(row.get('Disponibile', True)))

                        f_forza_salto = st.checkbox('✂️ Forza nuova pagina prima di questo piatto',
                                                    value=_safe_bool(row.get('Forza Salto Pagina', False)))

                        f_note = st.text_input('📝 Note Chef (non stampate)',
                                               value=_safe_str(row.get('Note Chef','')))

                        cb1, cb2 = st.columns(2)
                        with cb1:
                            salvato = st.form_submit_button('💾 Salva', type='primary',
                                                            use_container_width=True)
                        with cb2:
                            eliminato = st.form_submit_button('🗑️ Elimina', use_container_width=True)

                    if salvato:
                        for col_k, val_v in [
                            ('Nome IT', f_nome_it), ('Nome EN', f_nome_en),
                            ('Descrizione IT', f_desc_it), ('Descrizione EN', f_desc_en),
                            ('Prezzo', f_prezzo), ('Allergeni', f_allerg),
                            ('Visibile', f_vis), ('Disponibile', f_disp),
                            ('Forza Salto Pagina', f_forza_salto), ('Note Chef', f_note),
                        ]:
                            st.session_state.dati_menu.loc[sel, col_k] = val_v
                        st.success('✅ Salvato!')
                        st.rerun()

                    if eliminato:
                        st.session_state.dati_menu = (
                            st.session_state.dati_menu.drop(index=sel).reset_index(drop=True)
                        )
                        st.session_state.piatto_sel = None
                        st.rerun()

                # ── TAB LAYOUT ────────────────────────────────
                with tab_layout_ed:
                    with st.form(key=f'form_lay_{sel}'):
                        c1, c2 = st.columns(2)
                        with c1:
                            f_cat_it = st.text_input('📂 Categoria IT',
                                                     value=_safe_str(row.get('Categoria IT','')))
                        with c2:
                            f_cat_en = st.text_input('📂 Categoria EN',
                                                     value=_safe_str(row.get('Categoria EN','')))
                        c3, c4 = st.columns(2)
                        with c3:
                            f_colore = st.text_input('🎨 Colore categoria (hex)',
                                                     value=_safe_str(row.get('Colore Cat.','#b58d3d')))
                        with c4:
                            f_pagina = st.number_input('📄 Pagina logica', min_value=1,
                                                       value=_safe_int(row.get('Pagina',1)))
                        f_sep = st.selectbox('Separatore categoria',
                                             SEP_OPTIONS,
                                             index=SEP_OPTIONS.index(_norm_separatore(row.get('Separatore',SEP_LINEA))))
                        f_scala  = st.slider('🔎 Scala testo piatto', 0.4, 2.0,
                                             float(row.get('Scala Piatto', 1.0) or 1.0), 0.05)
                        f_spazio = st.slider('↕️ Spazio extra', 0.0, 5.0,
                                             float(row.get('Spazio Extra', 0.0) or 0.0), 0.25)

                        # Bottoni riordina
                        cu, cd, cs = st.columns(3)
                        with cu:
                            move_up   = st.form_submit_button('⬆️ Su', use_container_width=True)
                        with cd:
                            move_down = st.form_submit_button('⬇️ Giù', use_container_width=True)
                        with cs:
                            salva_lay = st.form_submit_button('💾 Salva', type='primary',
                                                              use_container_width=True)

                    if salva_lay:
                        for col_k, val_v in [
                            ('Categoria IT', f_cat_it), ('Categoria EN', f_cat_en),
                            ('Colore Cat.', f_colore), ('Pagina', int(f_pagina)),
                            ('Separatore', f_sep), ('Scala Piatto', f_scala),
                            ('Spazio Extra', f_spazio),
                        ]:
                            st.session_state.dati_menu.loc[sel, col_k] = val_v
                        st.success('✅ Layout salvato!')
                        st.rerun()

                    if move_up or move_down:
                        df_tmp = st.session_state.dati_menu.copy()
                        df_tmp = df_tmp.sort_values(['Pagina', 'Ordine']).reset_index()
                        pos_in_sorted = df_tmp[df_tmp['index'] == sel].index
                        if len(pos_in_sorted) > 0:
                            pos = pos_in_sorted[0]
                            if move_up and pos > 0:
                                other_orig = df_tmp.iloc[pos - 1]['index']
                                o1 = st.session_state.dati_menu.loc[sel, 'Ordine']
                                o2 = st.session_state.dati_menu.loc[other_orig, 'Ordine']
                                st.session_state.dati_menu.loc[sel, 'Ordine'] = o2
                                st.session_state.dati_menu.loc[other_orig, 'Ordine'] = o1
                            elif move_down and pos < len(df_tmp) - 1:
                                other_orig = df_tmp.iloc[pos + 1]['index']
                                o1 = st.session_state.dati_menu.loc[sel, 'Ordine']
                                o2 = st.session_state.dati_menu.loc[other_orig, 'Ordine']
                                st.session_state.dati_menu.loc[sel, 'Ordine'] = o2
                                st.session_state.dati_menu.loc[other_orig, 'Ordine'] = o1
                        st.rerun()

                # ── TAB AVANZATO: editor tabella completo ──────
                with tab_avanzato:
                    st.caption('Vista tabella completa — modifica diretta tutte le colonne.')
                    col_cfg_adv = {
                        'Ordine': st.column_config.NumberColumn('Ordine ↕', min_value=0, step=1, format='%d', width='small'),
                        'Pagina': st.column_config.NumberColumn('📄 Pag.', min_value=1, step=1, format='%d', width='small'),
                        'Separatore': st.column_config.SelectboxColumn('Sep.', options=SEP_OPTIONS, width='small'),
                        'Visibile':   st.column_config.CheckboxColumn('👁️', width='small'),
                        'Disponibile':st.column_config.CheckboxColumn('✅', width='small'),
                        'Forza Salto Pagina': st.column_config.CheckboxColumn('✂️', width='small'),
                        'Scala Piatto':st.column_config.NumberColumn('Scala', min_value=0.4, max_value=2.0, step=0.05, format='%.2f', width='small'),
                        'Spazio Extra':st.column_config.NumberColumn('Spazio', min_value=0.0, max_value=10.0, step=0.25, format='%.2f', width='small'),
                        'Categoria IT': st.column_config.TextColumn('Cat IT', width='medium'),
                        'Categoria EN': st.column_config.TextColumn('Cat EN', width='medium'),
                        'Colore Cat.':  st.column_config.TextColumn('🎨', width='small'),
                        'Nome IT':        st.column_config.TextColumn('Nome IT', width='large'),
                        'Nome EN':        st.column_config.TextColumn('Nome EN', width='large'),
                        'Descrizione IT': st.column_config.TextColumn('Desc IT', width='large'),
                        'Descrizione EN': st.column_config.TextColumn('Desc EN', width='large'),
                        'Prezzo':         st.column_config.TextColumn('€', width='small'),
                        'Allergeni':      st.column_config.TextColumn('Allerg.', width='small'),
                        'Note Chef':      st.column_config.TextColumn('Note', width='medium'),
                    }
                    df_adv_edit = st.data_editor(
                        st.session_state.dati_menu,
                        use_container_width=True, num_rows='dynamic', height=400,
                        column_config=col_cfg_adv,
                        column_order=['Ordine','Pagina','Separatore','Visibile','Disponibile','Forza Salto Pagina',
                                      'Scala Piatto','Spazio Extra','Categoria IT','Categoria EN',
                                      'Colore Cat.','Nome IT','Nome EN','Descrizione IT',
                                      'Descrizione EN','Prezzo','Allergeni','Note Chef'],
                        key='editor_tabella_adv',
                    )
                    if st.button('💾 Applica modifiche tabella', key='apply_adv', type='primary'):
                        st.session_state.dati_menu = _assicura_colonne_menu(df_adv_edit.copy())
                        st.success('✅ Tabella applicata.')
                        st.rerun()

        # ── Aggiorna sidebar con dati correnti ──────────────────
        df_cur = _assicura_colonne_menu(st.session_state.dati_menu.copy())
        with _mappa_placeholder:
            render_mappa_pagine(build_mappa_pagine(df_cur))
        with _layout_check_placeholder:
            render_layout_check(df_cur)

        # ── Anteprima e Esportazione ────────────────────────────
        st.divider()
        st.markdown('### 👁️ Anteprima e Esportazione')

        df_export = _assicura_colonne_menu(st.session_state.dati_menu.copy())
        base_font = round(zoom_foglio * 16.0, 2)

        html_content = genera_html(
            df=df_export,
            logo_b64=get_image_base64(logo_file),
            bg_b64=get_image_base64(bg_file),
            stile_sfondo=stile_sfondo,
            titolo_menu=titolo_menu,
            testo_footer=testo_footer,
            logo_size_px=logo_size,
            base_font_px=base_font,
            mostra_nome_en=mostra_nome_en,
            mostra_cat_en=mostra_cat_en,
            template_key=template_sel,
            disabilita_autopaginazione=disabilita_autopaginazione,
        )

        n_fogli = html_content.count('class="foglio-a4"')
        n_vis_e = int(df_export['Visibile'].apply(_safe_bool).sum())
        n_nd_e  = int((~df_export['Disponibile'].apply(_safe_bool)).sum())

        st.info(f'📄 Fogli: **{n_fogli}** · 👁️ {n_vis_e}/{len(df_export)} visibili · 🔴 {n_nd_e} non disponibili', icon='ℹ️')
        st.components.v1.html(html_content, height=1200 + (n_fogli - 1) * 1160, scrolling=True)

        st.divider()
        ce1, ce2, ce3, ce4 = st.columns(4)
        with ce1:
            st.download_button('💾 Progetto (.json)',
                               df_export.to_json(orient='records', force_ascii=False, indent=2),
                               f'{nome_file_menu}.json', 'application/json',
                               use_container_width=True)
        with ce2:
            try:
                st.download_button('📊 Excel',
                                   df_to_excel_bytes(df_export),
                                   f'{nome_file_menu}.xlsx',
                                   'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                   use_container_width=True)
            except Exception as e:
                st.warning(f'Excel: {e}')
        with ce3:
            st.download_button('📥 HTML Stampa',
                               html_content.encode('utf-8'),
                               f'{nome_file_menu}_Stampa.html',
                               'text/html; charset=utf-8',
                               use_container_width=True)
        with ce4:
            if PDF_DISPONIBILE:
                try:
                    with st.spinner('PDF…'):
                        pdf_b = WeasyHTML(string=html_content).write_pdf()
                    st.download_button('📄 PDF', pdf_b, f'{nome_file_menu}.pdf',
                                       'application/pdf', use_container_width=True)
                except Exception as e:
                    st.error(f'WeasyPrint: {e}')
            else:
                st.info('WeasyPrint non installato:\n```\npip install weasyprint\n```', icon='ℹ️')


# ═══════════════════════════════════════════════════════════════════
# TAB 2 — APERITIVI / TAGLIERI
# ═══════════════════════════════════════════════════════════════════

with tab_aperitivi:
    st.caption('Parser aperitivi/taglieri — sezione indipendente dal menu principale.')

    for k, v in [('dati_aperitivi', pd.DataFrame()), ('aper_json_id', None), ('aper_xlsx_id', None)]:
        if k not in st.session_state:
            st.session_state[k] = v

    st.markdown('### 📂 Caricamento Testi Aperitivi')
    ca1, ca2 = st.columns(2)
    with ca1:
        file_aper_it = st.file_uploader('🇮🇹 Word Aperitivi IT (.docx)', type=['docx'], key='aper_it')
    with ca2:
        file_aper_en = st.file_uploader('🇬🇧 Word Aperitivi EN (.docx) — opzionale', type=['docx'], key='aper_en')

    if st.button('🔄 Estrai Taglieri', type='primary', key='btn_aper'):
        if file_aper_it:
            with st.spinner('Parser aperitivi…'):
                righe_it = estrai_taglieri_word(file_aper_it)
                if file_aper_en:
                    en_map = {r['Nome IT']: r for r in estrai_taglieri_word(file_aper_en)}
                    for r in righe_it:
                        en_r = en_map.get(r['Nome IT'])
                        if en_r:
                            r['Descrizione EN'] = en_r.get('Descrizione IT', '')
                            r['Sottotitolo EN'] = en_r.get('Sottotitolo IT', '')
            st.session_state.dati_aperitivi = _assicura_colonne_aperitivi(pd.DataFrame(righe_it))
            st.session_state.aper_json_id = None; st.session_state.aper_xlsx_id = None
            n_tag  = sum(1 for r in righe_it if r['Tipo'] == 'Tagliere')
            n_voci = sum(1 for r in righe_it if r['Tipo'] == 'Voce')
            st.success(f'✅ Estratti **{n_tag}** taglieri e **{n_voci}** voci.')
        else:
            st.error('⚠️ Carica almeno il file Word italiano degli aperitivi.')

    if aper_progetto_json is not None:
        fid = id(aper_progetto_json)
        if fid != st.session_state.aper_json_id:
            try:
                st.session_state.dati_aperitivi = _assicura_colonne_aperitivi(pd.DataFrame(json.load(aper_progetto_json)))
                st.session_state.aper_json_id = fid
                st.success('✅ Progetto aperitivi JSON caricato.')
            except Exception as e:
                st.error(f'Errore JSON: {e}')

    if aper_progetto_xlsx is not None:
        fid = id(aper_progetto_xlsx)
        if fid != st.session_state.aper_xlsx_id:
            try:
                st.session_state.dati_aperitivi = _assicura_colonne_aperitivi(excel_bytes_to_df_aperitivi(aper_progetto_xlsx.getvalue()))
                st.session_state.aper_xlsx_id = fid
                st.success('✅ Dati Excel aperitivi caricati.')
            except Exception as e:
                st.error(f'Errore Excel: {e}')

    if not st.session_state.dati_aperitivi.empty:
        df_ab = _assicura_colonne_aperitivi(st.session_state.dati_aperitivi.copy())

        st.divider()
        st.markdown('### ✏️ Editor Aperitivi')

        na_tot = len(df_ab)
        na_vis = int(df_ab['Visibile'].apply(_safe_bool).sum())
        na_nd  = int((~df_ab['Disponibile'].apply(_safe_bool)).sum())
        na_tag = int((df_ab['Tipo'] == 'Tagliere').sum())
        na_voci= int((df_ab['Tipo'] == 'Voce').sum())
        ma1,ma2,ma3,ma4,ma5 = st.columns(5)
        ma1.metric('📋 Voci',        na_tot)
        ma2.metric('👁️ Visibili',     na_vis)
        ma3.metric('🔴 Non disp.',   na_nd)
        ma4.metric('🧀 Taglieri',    na_tag)
        ma5.metric('🍽️ Voci ingr.', na_voci)

        LAYOUT_OPTIONS_APER = [LAYOUT_NUOVA, LAYOUT_CONTINUA, LAYOUT_STESSA]
        col_cfg_a = {
            'Ordine':      st.column_config.NumberColumn('Ordine ↕', min_value=0, step=1, format='%d', width='small'),
            'Tipo':        st.column_config.SelectboxColumn('Tipo', options=['Tagliere','Voce'], width='small'),
            'Layout Pagina': st.column_config.SelectboxColumn('📄 Layout', options=LAYOUT_OPTIONS_APER, width='medium'),
            'Visibile':    st.column_config.CheckboxColumn('👁️', width='small'),
            'Disponibile': st.column_config.CheckboxColumn('✅', width='small'),
            'Scala':       st.column_config.NumberColumn('Scala', min_value=0.4, max_value=2.0, step=0.05, format='%.2f', width='small'),
            'Spazio Extra':st.column_config.NumberColumn('Spazio', min_value=0.0, max_value=10.0, step=0.25, format='%.2f', width='small'),
            'Nome IT':        st.column_config.TextColumn('Nome IT',       width='large'),
            'Nome EN':        st.column_config.TextColumn('Nome EN',       width='medium'),
            'Sottotitolo IT': st.column_config.TextColumn('Sottotitolo IT',width='large'),
            'Sottotitolo EN': st.column_config.TextColumn('Sottotitolo EN',width='large'),
            'Descrizione IT': st.column_config.TextColumn('Desc IT',       width='large'),
            'Descrizione EN': st.column_config.TextColumn('Desc EN',       width='large'),
            'Prezzo':         st.column_config.TextColumn('Prezzo',        width='small'),
            'Allergeni':      st.column_config.TextColumn('Allergeni',     width='medium'),
            'Note Chef':      st.column_config.TextColumn('Note Chef',     width='large'),
        }
        df_aper_mod = st.data_editor(
            df_ab, use_container_width=True, num_rows='dynamic', height=480,
            column_config=col_cfg_a,
            column_order=['Ordine','Tipo','Layout Pagina','Visibile','Disponibile',
                          'Scala','Spazio Extra','Nome IT','Nome EN','Sottotitolo IT','Sottotitolo EN',
                          'Descrizione IT','Descrizione EN','Prezzo','Allergeni','Note Chef'],
            key='editor_aperitivi',
        )

        st.divider()
        st.markdown('### 👁️ Anteprima e Esportazione Aperitivi')

        html_aper = genera_html_aperitivi(
            df=df_aper_mod, logo_b64=get_image_base64(logo_file),
            bg_b64=get_image_base64(bg_file), stile_sfondo=stile_sfondo,
            titolo_aperitivi=titolo_aperitivi, testo_footer=testo_footer,
            logo_size_px=logo_size, base_font_px=round(zoom_foglio*16.0,2),
            mostra_nome_en=mostra_nome_en, pagina_unica=pagina_unica_aper,
        )
        n_fogli_a = html_aper.count('class="foglio-a4"')
        st.components.v1.html(html_aper, height=1200 + (n_fogli_a-1)*1160, scrolling=True)
        st.info(f'📄 Fogli: **{n_fogli_a}** · 👁️ {na_vis}/{na_tot} · 🔴 {na_nd}', icon='ℹ️')

        st.divider()
        ae1,ae2,ae3,ae4 = st.columns(4)
        with ae1:
            st.download_button('💾 Progetto (.json)',
                               df_aper_mod.to_json(orient='records', force_ascii=False, indent=2),
                               f'{nome_file_aper}.json', 'application/json', use_container_width=True)
        with ae2:
            try:
                st.download_button('📊 Excel', df_to_excel_bytes_aperitivi(df_aper_mod),
                                   f'{nome_file_aper}.xlsx',
                                   'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                   use_container_width=True)
            except Exception as e:
                st.warning(f'Excel: {e}')
        with ae3:
            st.download_button('📥 HTML Stampa', html_aper.encode('utf-8'),
                               f'{nome_file_aper}_Stampa.html', 'text/html; charset=utf-8',
                               use_container_width=True)
        with ae4:
            if PDF_DISPONIBILE:
                try:
                    with st.spinner('PDF aperitivi…'):
                        pdf_aper = WeasyHTML(string=html_aper).write_pdf()
                    st.download_button('📄 PDF', pdf_aper, f'{nome_file_aper}.pdf',
                                       'application/pdf', use_container_width=True)
                except Exception as e:
                    st.error(f'WeasyPrint: {e}')
            else:
                st.info('WeasyPrint non installato:\n```\npip install weasyprint\n```', icon='ℹ️')
